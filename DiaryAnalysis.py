# coding: utf-8 
import docx
from snownlp import SnowNLP
import matplotlib.pyplot as plt

def cutsen(InputStringList,Divider):
    OutputStringList=[]
    for each in InputStringList:
        Temp=each.split(Divider)
        for every in Temp:
            OutputStringList.append(every)
    return OutputStringList

def cutintopieces():
    Temp=[]
    for i in range(1,7):
        file=docx.Document("D ("+str(i)+").docx")
        for i in range(len(file.paragraphs)):
            Sample=file.paragraphs[i].text
            Sample=Sample.split("，")
            Sample=cutsen(Sample,"？")
            Sample=cutsen(Sample,"。")
            for each in Sample:
                Temp.append(each)
    return Temp

def toEmo():
    emotion=[]
    Temp=cutintopieces()
    for each in Temp:
        if len(each)>0:
            s = SnowNLP(each)
            emotion.append(s.sentiments)
    return emotion

emotion=toEmo()
plt.plot(emotion)
